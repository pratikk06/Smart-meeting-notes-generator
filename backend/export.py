"""
Day 4 scope: export meeting notes as Word (.docx) or PDF, and save/list
meeting history so past meetings can be revisited.

Design notes:
- Export functions take the same `notes` dict shape produced by
  summarize_transcript(), so this module has no coupling to how the notes
  were generated (Gemini/Claude/etc. — doesn't matter, just the shape).
- We write to an in-memory buffer (BytesIO) rather than a temp file, so
  Streamlit's st.download_button can serve it directly without leaving
  files scattered on disk that need manual cleanup.
"""

import io
from docx import Document
from docx.shared import Pt
from fpdf import FPDF


def _notes_sections(notes: dict, transcript_text: str = None):
    """Shared structure walk so docx/pdf builders don't duplicate this logic."""
    sections = [("Summary", [notes.get("summary", "")])]

    decisions = notes.get("decisions", [])
    sections.append(("Decisions", decisions if decisions else ["No decisions detected."]))

    action_items = notes.get("action_items", [])
    if action_items:
        lines = [
            f"{item.get('task', '')} — Owner: {item.get('owner') or 'Unassigned'}, "
            f"Deadline: {item.get('deadline') or 'Not specified'}"
            for item in action_items
        ]
    else:
        lines = ["No action items detected."]
    sections.append(("Action Items", lines))

    open_qs = notes.get("open_questions", [])
    sections.append(("Open Questions", open_qs if open_qs else ["No open questions detected."]))

    if transcript_text:
        sections.append(("Full Transcript", [transcript_text]))

    return sections


def export_to_docx(notes: dict, meeting_title: str = "Meeting Notes",
                    transcript_text: str = None) -> bytes:
    """Build a .docx file in memory and return its raw bytes."""
    doc = Document()
    doc.add_heading(meeting_title, level=0)

    for heading, items in _notes_sections(notes, transcript_text):
        doc.add_heading(heading, level=1)
        for item in items:
            p = doc.add_paragraph(item, style="List Bullet" if len(items) > 1 else None)
            p.paragraph_format.space_after = Pt(4)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_to_pdf(notes: dict, meeting_title: str = "Meeting Notes",
                   transcript_text: str = None) -> bytes:
    """Build a .pdf file in memory and return its raw bytes."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    pdf.set_font("Helvetica", "B", 18)
    pdf.multi_cell(0, 10, meeting_title)
    pdf.set_x(pdf.l_margin)
    pdf.ln(2)

    for heading, items in _notes_sections(notes, transcript_text):
        pdf.set_font("Helvetica", "B", 13)
        pdf.multi_cell(0, 8, heading)
        pdf.set_x(pdf.l_margin)
        pdf.set_font("Helvetica", "", 11)
        for item in items:
            # encode/decode strips characters FPDF's core font can't render
            safe_text = item.encode("latin-1", "replace").decode("latin-1")
            pdf.multi_cell(0, 6, f"- {safe_text}")
            pdf.set_x(pdf.l_margin)
        pdf.ln(3)

    # fpdf2 returns a bytearray; Streamlit's download_button wants bytes
    return bytes(pdf.output())